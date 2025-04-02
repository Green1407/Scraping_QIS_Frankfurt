"""
This module handles the creation, modification, and export of the Word documents.
It inserts lecturer information, semester details, and course data from the scraped course data dataframe
into structured tables while ensuring proper formatting and alignment.
"""

import os

from docx import Document
from docx.shared import Pt
from docx.table import _Cell
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pandas as pd

def fill_word(name: str, semester: str, data: list[str], institute_data: pd.DataFrame) -> Document:
    """
    Fills the Word document template with general information about the lecturer and his course data.

    Parameters:
        name (str): The name of the person for whom the document is generated.
        semester (str): The semester information.
        data (list[str]): A list of course data.
        institute_data (pd.DataFrame): A Dataframe with the lecturers institute data

    Returns:
        Document: The modified result Word document.
    """
    # Select the appropriate template based on the number of courses, select the bigger one if more than 9 courses
    link = "Beispiel_big.docx" if len(data) > 9 else "Beispiel.docx"

    # If there is no data do not export as word doc and alternativly do not export if not desired institute
    if len(data) < 1: #or get_institute(name,institute_data) != "Institut für Informatik (IfI):
        return 0

    # Warning if lecturer is responsible for more than 18 courses
    if len(data) > 18:
        print(f"Achtung für {name} im Semester {semester} wurde die Maximalanzahl von 18 Veranstaltungen überschritten.")

    # Load the Word document
    doc = read_word(link)

    # Write general information (name and semester)
    doc = write_general_info(doc,name,semester,institute_data)

    # Write course data
    doc = write_courses(doc,data)

    # Make sure the export directory exists and export the new doc with an appropriate title
    os.makedirs("Word_Exporte", exist_ok=True)
    doc.save("Word_Exporte/" + semester.replace(" ", "").replace("/","-") + "_" + name.replace(" ", "_").replace("*","").replace("/","-") + "_Formular_A38.docx")

    return doc

def read_word(link: str) -> Document:
    """
    Reads a Word document file and returns the document object.

    Parameter:
        link (str): The file path to the Word document.

    Returns:
        Document: A docx.Document object representing the loaded Word document.
    """
    doc = Document(link)
    return(doc)

def write_courses(doc: Document, data: pd.DataFrame) -> Document:
    """
    Writes the course information from the given data into the given Word document.

    Parameters:
        doc (Document): The Word document object where the course information will be written.
        data (pd.DataFrame): A DataFrame containing course data with the following expected columns:
            - 'Kursname' (str): The course name.
            - 'Veranstaltungsart' (str): The type of course.
            - 'SWS' (int/float): The number of semester hours.
            - 'Lehrpersonen' (List[str]): A list of other lecturers.

    Returns:
        Document: The modified Word document with course information inserted.
    """
    # Starting index for insertion into the word table, this number need to be changed when using another word template
    table_index = 20

    #Iterate through all data in the dataframe, get data for each course and write it in the Word doc
    for row in data.itertuples(index=False):
        course_name = row.Kursname
        course_type = row.Veranstaltungsart
        course_sws = row.SWS

        #If more than 2 other persons are involved only show the first 2 in the export doc
        course_persons = ', '.join(row.Lehrpersonen[:2]) + (
            ' und weitere' if len(row.Lehrpersonen) > 2 else '') if row.Lehrpersonen else '-'

        # Course data is written in to the doc here
        doc = write_course(table_index,doc,course_name,course_type,course_sws,course_persons,course_sws)

        #Increment index for the next course, increment is by 2 because in the sample file each cell has 2 rows
        table_index = table_index + 2

    return doc

def write_general_info(doc: Document, name: str, semester: str, institute_data: pd.DataFrame) -> Document:
    """
    Writes the general information (name and semester) into the first table of the Word document.

    Parameters:
        doc (Document): The sample Word Document that will be filled.
        name (str): The lecturers name that will be inserted at the top of the document.
        semester (str): The semester information to be inserted also at the top of the document.
        institute_dataframe (pd.DataFrame) : DataFrame with the corresponding data of the lecturers institutes

    Returns:
        Document: The modified Word document object.
    """
    table = doc.tables[0]

    # Write name
    write_text(table.cell(3, 3), f"Name: {name}", font_size=12)

    # Write semester
    write_text(table.cell(5, 3), f"Semester: {semester}", font_size=12)

    # Write institute
    write_text(table.cell(5, 15), f"Institut: {get_institute(name,institute_data)}", font_size=12)

    return(doc)

def get_institute(name: str, institute_dataframe: pd.DataFrame) -> str:
    """
    Takes a name of a lecturer as input, then search through the personal and institute list in the Database folder and
    returns the corresponding institute as string

    Parameters:
        name (str) : name string of the lecturer
        institute_dataframe (pd.DataFrame) : DataFrame with the corresponding data of the lecturers institutes

    Returns:
        institute (str) or str: institute of the lecturer as string or Bitte manuell einfügen
      """
    institute = institute_dataframe.loc[institute_dataframe['Person'] == name, 'Institut']

    return institute.iloc[0] if not institute.empty else "Bitte manuell einfügen"

def write_course(
    table_index: int,
    doc: Document,
    name: str,
    course_type: str,
    sws: int,
    other_teachers: str,
    fullfilled_sws: int
) -> Document:
    """
    Writes a single course entry into the appropriate table row in the Word document.

    Parameters:
        table_index (int): The row index in the table where the course information should be inserted.
        doc (Document): The Word document object containing the tables.
        name (str): The name of the course.
        course_type (str): The type of course (e.g., lecture, seminar).
        sws (int): The semester weekly hours (SWS) assigned to the course.
        other_teachers (str): The names of additional teachers.
        fulfilled_sws (int): The number of fulfilled semester weekly hours.

    Returns:
        Document: The modified Word document with the course information added.
    """
    table = doc.tables[0] # Start with the first table

    # Check if the table index is beyond the available rows
    if table_index >= len(table.rows):
        table = doc.tables[1] # If beyond: Switch to the second table
        table_index = table_index - 33 # Adjust index for the second table

        # If the adjusted index is still out of bounds, return the document unchanged, extremely unlikely to happen
        if table_index >= len(table.rows):
            return(doc)

    # Insert course details into the correct cells, using specified formatting
    write_text(table.cell(table_index, 8), str(name),font_name="Arial Narrow", font_size=8, italic=True, align="center")
    write_text(table.cell(table_index, 12), str(course_type),font_name="Arial Narrow", font_size=8, italic=True, align="center")
    write_text(table.cell(table_index, 16), str(sws),font_name="Arial Narrow", font_size=8, italic=True, align="center")
    write_text(table.cell(table_index, 20), str(other_teachers),font_name="Arial Narrow", font_size=8, italic=True, align="center")
    write_text(table.cell(table_index, 24), str(fullfilled_sws),font_name="Arial Narrow", font_size=8, italic=True, align="center")

    return doc


def write_text(
        cell: _Cell,
        text: str,
        font_name: str = "Arial",
        font_size: int = 12,
        italic: bool = False,
        align: str = "left"
) -> None:
    """
    Writes formatted text into a Word table cell with given font name and size

    Parameters:
        cell (_Cell): The table cell where the text will be inserted.
        text (str): The text to write in the cell.
        font_name (str, optional): The font name to use (default: "Arial").
        font_size (int, optional): The font size in points (default: 12).
        italic (bool, optional): Whether the text should be italicized (default: False).
        align (str, optional): Text alignment ("left", "center", or "right", default: "left").

    Returns:
        None
    """
    # Clear existing text
    cell.text = ""

    # Add a new paragraph and run
    paragraph = cell.add_paragraph()
    run = paragraph.add_run(text)

    # Set font properties
    run.font.name = font_name
    run.font.size = Pt(font_size)

    # Set italic style if specified
    run.italic = italic

    # Set text alignment
    if align == "center":
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif align == "right":
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    else:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  #Default setup is left
